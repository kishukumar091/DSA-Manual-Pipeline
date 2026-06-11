"""
builder/sections.py — Renders one problem section into the Word document.

Each section contains:
  1. Coloured heading badge + problem title
  2. Shaded monospace code table
  3. Verification screenshot (if available)
  4. Italic caption
"""
from docx import Document
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config
from problems.registry import ProblemEntry


# ─────────────────────────────────────────────────────────────────────────────
# XML helpers
# ─────────────────────────────────────────────────────────────────────────────
def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply solid background fill to a table cell via raw XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = parse_xml(
        f'<w:shd {nsdecls("w")} '
        f'w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    tcPr.append(shd)


def _set_para_shading(para, hex_color: str) -> None:
    """Apply paragraph-level shading (used for heading badge background)."""
    pPr = para._p.get_or_add_pPr()
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} '
        f'w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    pPr.append(shd)


def _remove_table_borders(table) -> None:
    """Remove all visible borders from a table (borderless panel look)."""
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top    w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left   w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right  w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))


def _set_table_width(table, width_twips: int) -> None:
    """Force a table to an absolute width (in twentieths of a point)."""
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = parse_xml(
        f'<w:tblW {nsdecls("w")} '
        f'w:w="{width_twips}" w:type="dxa"/>'
    )
    tblPr.append(tblW)


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────
def add_category_divider(doc: Document, category: str) -> None:
    """
    Full-width dark banner that marks the start of a new category group.
    Example:  ── Arrays & Hashing ──
    """
    hex_color = config.CATEGORY_COLORS.get(category, config.CATEGORY_COLORS["Other"])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_shading(p, hex_color)

    run = p.add_run(f"  {'─' * 6}  {category.upper()}  {'─' * 6}  ")
    run.bold             = True
    run.font.name        = config.FONT_HEAD
    run.font.size        = Pt(11)
    run.font.color.rgb   = RGBColor(0xFF, 0xFF, 0xFF)


def add_problem_section(doc: Document,
                        entry: ProblemEntry,
                        image_path: str = None) -> None:
    """Render a full section for one problem."""

    # 1. ── Heading ────────────────────────────────────────────────────────────
    hex_color = config.CATEGORY_COLORS.get(entry.category,
                                            config.CATEGORY_COLORS["Other"])

    heading_p = doc.add_paragraph()
    heading_p.paragraph_format.keep_with_next = True
    heading_p.paragraph_format.space_before   = Pt(14)
    heading_p.paragraph_format.space_after    = Pt(0)

    # Programme badge:  [Prog 1]
    badge = heading_p.add_run(f"  Prog {entry.number:02d}  ")
    badge.bold           = True
    badge.font.name      = config.FONT_HEAD
    badge.font.size      = config.SIZE_H3
    badge.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Inline character highlight colour
    rPr = badge._r.get_or_add_rPr()
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} '
        f'w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    rPr.append(shd)

    # Title text
    title_run = heading_p.add_run(f"  {entry.title}")
    title_run.bold           = True
    title_run.font.name      = config.FONT_HEAD
    title_run.font.size      = config.SIZE_H3
    title_run.font.color.rgb = config.COL_PRIMARY

    # Session label on same line
    tab_run = heading_p.add_run(f"  |  Session {entry.session}")
    tab_run.font.name      = config.FONT_BODY
    tab_run.font.size      = Pt(9)
    tab_run.font.color.rgb = config.COL_MUTED
    tab_run.italic         = True

    # Underline rule beneath heading
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(2)
    rule_p.paragraph_format.space_after  = Pt(6)
    rule_run = rule_p.add_run("─" * 90)
    rule_run.font.color.rgb = config.COL_MUTED
    rule_run.font.size      = Pt(6)

    # 2. ── Code block ─────────────────────────────────────────────────────────
    table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    _set_table_width(table, 9350)  # ~6.5" in twips (1" = 1440 twips)

    cell = table.cell(0, 0)
    _set_cell_shading(cell, config.COL_CODE_BG)

    # Cell margins (padding)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top    w:w="100" w:type="dxa"/>'
        f'  <w:left   w:w="180" w:type="dxa"/>'
        f'  <w:bottom w:w="100" w:type="dxa"/>'
        f'  <w:right  w:w="180" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

    # Write each line of code as a separate paragraph to preserve line breaks
    code_lines = entry.code.split("\n")
    for idx, line in enumerate(code_lines):
        if idx == 0:
            code_p = cell.paragraphs[0]
        else:
            code_p = cell.add_paragraph()
        code_p.paragraph_format.space_before = Pt(0)
        code_p.paragraph_format.space_after  = Pt(0)
        run = code_p.add_run(line if line else " ")
        run.font.name      = config.FONT_CODE
        run.font.size      = config.SIZE_CODE
        run.font.color.rgb = config.COL_CODE_FG

    doc.add_paragraph()  # spacer

    # 3. ── Screenshot (optional) ──────────────────────────────────────────────
    if image_path:
        img_p = doc.add_paragraph()
        img_p.alignment                    = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_before = Pt(6)
        img_p.paragraph_format.space_after  = Pt(4)
        img_p.paragraph_format.keep_with_next = True

        try:
            img_p.add_run().add_picture(image_path, width=config.IMG_WIDTH)
        except Exception as e:
            err_run = img_p.add_run(f"[Image unavailable: {e}]")
            err_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            err_run.italic = True

        # Caption
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(0)
        cap_p.paragraph_format.space_after  = Pt(18)
        c_run = cap_p.add_run(
            f"Figure {entry.number}: Verification screenshot - {entry.title}"
        )
        c_run.italic         = True
        c_run.font.size      = config.SIZE_CAP
        c_run.font.color.rgb = config.COL_CAP
    else:
        # placeholder note
        note_p = doc.add_paragraph()
        note_p.paragraph_format.space_before = Pt(4)
        note_p.paragraph_format.space_after  = Pt(18)
        note_run = note_p.add_run("[ No verification screenshot attached ]")
        note_run.italic         = True
        note_run.font.size      = Pt(9)
        note_run.font.color.rgb = config.COL_MUTED

    # Page break after every 2nd program to keep layout tidy (optional heuristic)
    # Uncomment the line below if you want hard page breaks between entries:
    # doc.add_page_break()
