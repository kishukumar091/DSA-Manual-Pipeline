"""
builder/document.py — Assembles the complete A4 Word document.

Pipeline:
  1. Page setup (A4, 1" margins)
  2. Global Normal style
  3. Title banner
  4. Table of Contents (static, text-based)
  5. Category dividers + per-problem sections
"""
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config
from problems.registry import PROBLEMS, ProblemEntry
from builder.sections import add_problem_section, add_category_divider


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────
def _set_para_shading(para, hex_color: str) -> None:
    pPr = para._p.get_or_add_pPr()
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} '
        f'w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    pPr.append(shd)


def _add_title_banner(doc: Document) -> None:
    """Dark full-width title banner with subtitle."""
    # Main title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(0)
    _set_para_shading(title_p, config.COL_DIV_BG)

    t1 = title_p.add_run("  DSA Reference Manual  ")
    t1.bold             = True
    t1.font.name        = config.FONT_HEAD
    t1.font.size        = config.SIZE_H1
    t1.font.color.rgb   = config.COL_DIV_FG

    # Subtitle line
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after  = Pt(24)
    _set_para_shading(sub_p, "1E293B")

    s1 = sub_p.add_run("  39 C++ Programs  ·  Lab Assignment Submissions  ")
    s1.font.name      = config.FONT_BODY
    s1.font.size      = Pt(11)
    s1.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)  # slate-400


def _add_toc(doc: Document, problems: list) -> None:
    """Static plain-text table of contents."""
    toc_heading = doc.add_paragraph()
    toc_heading.paragraph_format.space_before = Pt(0)
    toc_heading.paragraph_format.space_after  = Pt(8)
    h = toc_heading.add_run("Contents")
    h.bold             = True
    h.font.name        = config.FONT_HEAD
    h.font.size        = config.SIZE_H2
    h.font.color.rgb   = config.COL_PRIMARY

    # Rule
    rule = doc.add_paragraph("─" * 90)
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after  = Pt(4)
    rule.runs[0].font.size      = Pt(6)
    rule.runs[0].font.color.rgb = config.COL_MUTED

    # Two-column simulation using tab stops
    for entry in problems:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)

        num_run = p.add_run(f"  {entry.num:02d}.")
        num_run.bold           = True
        num_run.font.name      = config.FONT_BODY
        num_run.font.size      = Pt(9)
        num_run.font.color.rgb = config.COL_ACCENT

        title_run = p.add_run(f"  {entry.title}")
        title_run.font.name      = config.FONT_BODY
        title_run.font.size      = Pt(9)
        title_run.font.color.rgb = config.COL_TEXT

        tag_run = p.add_run(f"  [{entry.category}]")
        tag_run.italic         = True
        tag_run.font.name      = config.FONT_BODY
        tag_run.font.size      = Pt(8)
        tag_run.font.color.rgb = config.COL_MUTED

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# Main builder
# ─────────────────────────────────────────────────────────────────────────────
def build_document(
    image_map: dict,
    output_path: str,
    problem_numbers: list = None,
) -> None:
    """
    Build and save the DSA Reference Manual.

    Args:
        image_map      : {problem_num: abs_image_path}  (may be empty)
        output_path    : where to save the .docx file
        problem_numbers: optional list of ints to include; None = all 39
    """
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width   = config.PAGE_WIDTH
    section.page_height  = config.PAGE_HEIGHT
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, config.MARGIN)

    # ── Global styles ────────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name      = config.FONT_BODY
    normal.font.size      = config.SIZE_BODY
    normal.font.color.rgb = config.COL_TEXT

    # ── Title banner ─────────────────────────────────────────────────────────
    _add_title_banner(doc)

    # ── Filter + sort problems ────────────────────────────────────────────────
    entries = sorted(PROBLEMS.values(), key=lambda e: e.num)
    if problem_numbers:
        entries = [e for e in entries if e.num in problem_numbers]

    # ── Table of contents ────────────────────────────────────────────────────
    _add_toc(doc, entries)

    # ── Problem sections ──────────────────────────────────────────────────────
    prev_category = None
    for entry in entries:
        # Inject category divider when category changes
        if entry.category != prev_category:
            add_category_divider(doc, entry.category)
            prev_category = entry.category

        image_path = image_map.get(entry.num)
        add_problem_section(doc, entry, image_path)

    # ── Save ─────────────────────────────────────────────────────────────────
    doc.save(output_path)
    print(f"\n✓  Document saved → {output_path}")
    print(f"   Problems   : {len(entries)}")
    print(f"   With images: {sum(1 for e in entries if e.num in image_map)}")
